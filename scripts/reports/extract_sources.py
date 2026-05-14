from __future__ import annotations

import json
import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "p": "http://schemas.openxmlformats.org/presentationml/2006/main",
}


def clean(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def pdf_text(path: Path) -> dict:
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        pages.append({"page": i, "text": clean(page.extract_text() or "")})
    return {"type": "pdf", "pages": pages, "page_count": len(pages)}


def docx_text(path: Path) -> dict:
    doc = Document(str(path))
    paragraphs = [clean(p.text) for p in doc.paragraphs if clean(p.text)]
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([clean(cell.text) for cell in row.cells])
        if rows:
            tables.append(rows)
    return {
        "type": "docx",
        "paragraphs": paragraphs,
        "tables": tables,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
    }


def pptx_text(path: Path) -> dict:
    slides = []
    with zipfile.ZipFile(path) as zf:
        slide_names = sorted(
            [n for n in zf.namelist() if re.match(r"ppt/slides/slide\d+\.xml$", n)],
            key=lambda n: int(re.search(r"slide(\d+)\.xml$", n).group(1)),
        )
        for idx, name in enumerate(slide_names, start=1):
            root = ET.fromstring(zf.read(name))
            texts = []
            for node in root.findall(".//a:t", NS):
                if node.text and clean(node.text):
                    texts.append(clean(node.text))
            slide_text = clean("\n".join(texts))
            slides.append({"slide": idx, "text": slide_text})
    return {"type": "pptx", "slides": slides, "slide_count": len(slides)}


def value_to_text(value) -> str:
    if value is None:
        return ""
    return clean(str(value))


def xlsx_profile(path: Path) -> dict:
    wb = load_workbook(path, data_only=False, read_only=True)
    sheets = []
    for ws in wb.worksheets:
        non_empty = 0
        sample_rows = []
        header_candidate = None
        for row in ws.iter_rows():
            values = [value_to_text(cell.value) for cell in row]
            if any(values):
                non_empty += sum(1 for v in values if v)
                compact = values[:20]
                if header_candidate is None:
                    header_candidate = compact
                if len(sample_rows) < 12:
                    sample_rows.append(compact)
        sheets.append(
            {
                "name": ws.title,
                "max_row": ws.max_row,
                "max_column": ws.max_column,
                "non_empty_cells": non_empty,
                "header_candidate": header_candidate or [],
                "sample_rows": sample_rows,
            }
        )
    return {"type": "xlsx", "sheets": sheets, "sheet_count": len(sheets)}


def extract(path: Path) -> dict:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        payload = pdf_text(path)
    elif suffix == ".docx":
        payload = docx_text(path)
    elif suffix in {".pptx", ".ppt"}:
        payload = pptx_text(path)
    elif suffix in {".xlsx", ".xlsm"}:
        payload = xlsx_profile(path)
    else:
        payload = {"type": suffix.lstrip("."), "note": "unsupported"}
    return {"file": path.name, **payload}


def main() -> None:
    root = Path(sys.argv[1]).resolve()
    out = Path(sys.argv[2]).resolve()
    files = sorted(
        [
            p
            for p in root.iterdir()
            if p.is_file()
            and p.suffix.lower() in {".pdf", ".docx", ".pptx", ".ppt", ".xlsx", ".xlsm"}
        ],
        key=lambda p: p.name.lower(),
    )
    data = [extract(path) for path in files]
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {out} with {len(data)} source files")


if __name__ == "__main__":
    main()
